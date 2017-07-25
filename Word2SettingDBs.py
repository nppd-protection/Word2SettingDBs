#! /usr/bin/python

'''
USAGE:
Word2SettingDBs.py <Setting Calcs.docx> <AspenDB Export Settings.txt>
<Setting Calcs.docx>:  MS Word docx setting calculation document
<AspenDB Export Settings.txt>: Exported setting sheet from Aspen database

The order of the filenames can be transposed as long as one file has a docx
extension to identify it as the MS Word document.

ABOUT:
The program will create two files in the same directory as
<Setting Calcs.docx>.
Assuming the input docx filename is "Substation Setting Calcs.docx", the files
created will be as follows:
Substation Setting Calcs FOR ACSELERATOR.txt:  Text file for import into
                                               AcSELerator.
Substation Setting Calcs FOR ASPENDB.txt:  Text file for import into Aspen
                                           database.

This program is written to assist in getting relay settings developed in MS
Word calculation documents entered into the AcSELerator and Aspen databases.

IMPORTANT:  This script is intended to save time, but the protection engineer
is still responsible to ensure that settings get into the databases correctly.
Many factors could cause settings not to be transferred correctly, so thorough
checking continues to be an important part of the process.

The settings are extracted from the MS Word .docx file using style information,
so it is necessary to utilize the appropriate style in the calculation document
for the settings to be recognized.  The file must be saved in docx format or
this program will not be able to read it.

The AcSELerator file is written out based solely on the setting and group
information extracted from the docx setting calculation document.  All the
setting groups are written to one output text file that can be imported into
AcSELerator.  As of 7/14/2014, AcSELerator will merge in settings from a text
file even with no relay model info in the text file and without all settings
being present.  At this time it does have a bug where it clears the Event
Report Digitals even if these settings are not present in the imported text
file.  The user must unselect the import of this setting to avoid having the
Event Report Digitals cleared in the open setting sheet.

The Aspen Database Text File is not written from scratch, rather, an Aspen Text
File must be exported from the Aspen Database first, then this program reads in
that file, changes the setting values that it found in the docx setting
calculation document, then writes out a new Aspen Database text file for
subsequent import into Aspen Database.

Progress information and errors are logged to the same directory as the program
is run from.

'''

from __future__ import print_function, unicode_literals

import re
import sys
import os.path
import os
import codecs
import msvcrt

# Set up a logger so any errors can go to file to facilitate debugging
import logging
from logging.config import dictConfig


#
# The docx package documentation can be found at
# https://python-docx.readthedocs.org/en/latest/
#
# Installation:
# If pip is available:
# pip install python-docx
#
# If easy_install is available:
# easy_install python-docx
#
# Otherwise download the tar.gz and run setup.py install.
# Currently working with 0.8.5. Version 0.7.2 NOT working.
#
from docx import Document

# By default, log to the same directory the program is run from
if os.path.exists(os.path.dirname(sys.argv[0])):
    logfile = os.path.join(os.path.dirname(sys.argv[0]), 'Word2SettingDBs.log')
else:
    logfile = 'Word2SettingDBs.log'

logging_config = {
    'version': 1,
    'formatters': {
        'file': {'format':
                 '%(asctime)s ' + os.environ['USERNAME'] +
                 ' %(levelname)-8s %(message)s'},
        'console': {'format':
                    '%(levelname)-8s %(message)s'}
        },

    'handlers': {
        'file': {'class': 'logging.FileHandler',
                 'filename': logfile,
                 'formatter': 'file',
                 'level': 'INFO'},
        'console': {'class': 'logging.StreamHandler',
                    'formatter': 'console',
                    'level': 'INFO'}
        },
    'loggers': {
        'root': {'handlers': ['file', 'console'],
                 'level': 'DEBUG'}
        }
}

dictConfig(logging_config)

logger = logging.getLogger('root')

try:
    logger.info('Running %s.' % sys.argv[0])
    logger.info('Logging to file %s.' % os.path.abspath(logfile))

    debug = False
    # Fix issue with output encoding of special characters on Windows terminal
    # for Python 2.7 only.
    if sys.version_info[0] < 3:
        sys.stdout = codecs.getwriter(sys.stdout.encoding)(sys.stdout,
                                                           errors='replace')

    if len(sys.argv) < 3:
        logger.error("Not enough input parameters.  Please include two "
                     "filenames when calling this program.")
        logger.error(__doc__)
        raise SystemExit
    elif len(sys.argv) > 3:
        logger.error("Too many input parameters.  Please include two filenames"
                     " when calling this program.")
        logger.error(__doc__)
        raise SystemExit

    if re.match('.*\.docx$', sys.argv[1], flags=re.I):
        documentParam = sys.argv[1]
        aspenTemplateParam = sys.argv[2]
    else:
        documentParam = sys.argv[2]
        aspenTemplateParam = sys.argv[1]

    calc_file_base = re.match('(.*)\.doc[xm]$', documentParam, flags=re.I)
    sel_save_file = calc_file_base.group(1) + ' FOR ACSELERATOR.txt'
    aspen_save_file = calc_file_base.group(1) + ' FOR ASPEN.txt'
    logger.info('Input files: Word = ' + documentParam + ', Aspen template = '
                + aspenTemplateParam)
    logger.info('Output files: ' + sel_save_file + ', ' + aspen_save_file)

    def partition(pred, iterable):
        '''Divide a list into two lists based on a filtering function.
        Takes a function returning a boolean and an iterable (e.g. a list).
        Returns two lists, the first where the filtering function returned
        True, the second where the filtering function returned false.
        '''
        trues = []
        falses = []
        for item in iterable:
            if pred(item):
                trues.append(item)
            else:
                falses.append(item)
        return trues, falses

    def setting_run(r):
        return r.font.bold and not r.font.strike

    def invert_dict(d):
        return dict(map(lambda x: (x[1], x[0]), d.items()))

    def stripall(l, c=None):
        return list(map(lambda s: s.strip(c), l))

    def findTable(tables, func):
        '''
        Find the first table with cell text returning true when passed to func
        Most likely func should be a regex.match function.
        '''
        for t in tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        if func(p.text):
                            return t

    def findSetting(setting_list, setting):
        '''
        Find the first entry in setting_list with setting for the setting name.
        Returns setting value from setting_list. If not found, it returns None.
        '''
        for entry in setting_list:
            if entry[0] == setting:
                return entry[1]
        return None

    # A regular expression to find special characters.
    reSpecial = re.compile('[^-_ A-Za-z0-9]+')

    def portX(t):
        return "PORT "+re.match("PORT ([F,1-5]).*", t, flags=re.I).group(1)

    reGroups = [("ALIAS", re.compile("ALIAS", flags=re.I), ),
                ("GLOBAL", re.compile("GLOBAL", flags=re.I), ),
                ("PORT 87", re.compile("PORT 87", flags=re.I)),
                ("CHANNEL X", re.compile("CHANNEL X", flags=re.I)),
                ("CHANNEL Y", re.compile("CHANNEL Y", flags=re.I)),
                ("BREAKER MONITOR", re.compile("BREAKER MONITOR", flags=re.I)),
                ("ZONE 1",  re.compile("ZONE 1", flags=re.I)),
                ("GROUP 1",  re.compile("GROUP 1", flags=re.I)),
                ("GROUP 1 LOGIC",  re.compile("LOGIC", flags=re.I)),
                ("PROTECTION LOGIC", re.compile("PROTECTION LOGIC",
                                                flags=re.I)),
                ("AUTOMATION LOGIC", re.compile("AUTOMATION", flags=re.I)),
                ("OUTPUT", re.compile("OUTPUT", flags=re.I)),
                ("FRONT PANEL", re.compile("FRONT PANEL", flags=re.I)),
                ("REPORT",  re.compile("REPORT", flags=re.I)),
                ("TEXT",  re.compile("TEXT", flags=re.I)),
                # portX must be last.
                (portX,  re.compile("PORT ([F,1-5])", flags=re.I))]

    sel4XX_names = {"ALIAS": "T1",
                    "GLOBAL": "G1",
                    "PORT 87": "P87", # SEL-411L
                    "BREAKER MONITOR": "SM",
                    "MONITOR": "M1",  # SEL-487V
                    "ZONE 1": "Z1",   # SEL-487B
                    "GROUP 1": "S1",
                    "PROTECTION LOGIC": "L1",
                    "AUTOMATION LOGIC": "A1",
                    "OUTPUT": "O1",
                    "FRONT PANEL": "F1",
                    "REPORT": "R1",
                    "PORT 1": "P1",
                    "PORT 2": "P2",
                    "PORT 3": "P3",
                    "PORT 5": "P5",
                    "PORT F": "PF"}

    sel4XX_AspenDBnames = {"GLOBAL": "1GLOB",
                           "PORT 87": "1PORT87",
                           "BREAKER MONITOR": "2BRKR",
                           "MONITOR": "2MON",
                           "GROUP 1": "3GRP",
                           "PROTECTION LOGIC": "4PROT",
                           "AUTOMATION LOGIC": "5AUTO",
                           "OUTPUT": "6OUT",
                           "FRONT PANEL": "7FRNT",
                           "REPORT": "8RPT",
                           "PORT 1": "9PRT1",
                           "PORT 2": "9PRT2",
                           "PORT 3": "9PRT3",
                           "PORT 5": "9PRT5",
                           "PORT F": "9PRTF"}

    sel487B_AspenDBnames = {"ALIAS": "1ALIAS",
                            "GLOBAL": "2GLOB",
                            "ZONE 1": "3GRP",
                            "GROUP 1": "3GRP",
                            "PROTECTION LOGIC": "4PROT",
                            "AUTOMATION LOGIC": "5AUTO",
                            "OUTPUT": "6OUT",
                            "FRONT PANEL": "7FRNT",
                            "REPORT": "8RPT",
                            "PORT 1": "9PRT1",
                            "PORT 2": "9PRT2",
                            "PORT 3": "9PRT3",
                            "PORT 5": "9PRT5",
                            "PORT F": "9PRTF"}
    sel3XX_names = {"GLOBAL": "G",
                    "GROUP 1": "1",
                    "GROUP 1 LOGIC": "L1",
                    "REPORT": "R",
                    "CHANNEL X": "X",
                    "CHANNEL Y": "Y",
                    "TEXT": "T",
                    "PORT 1": "P1",
                    "PORT 2": "P2",
                    "PORT 3": "P3",
                    "PORT 4": "P4",
                    "PORT 5": "P5",
                    "PORT F": "PF"}

    sel3XX_AspenDBnames = {"GLOBAL": "1GLOB",
                           "GROUP 1": "2GRP",
                           "GROUP 1 LOGIC": "3LOGIC",
                           "REPORT": "4RPT",
                           "CHANNEL X": "6CHX",
                           "CHANNEL Y": "6CHY",
                           "TEXT": "5TEXT",
                           "PORT 1": "9PRT1",
                           "PORT 2": "9PRT2",
                           "PORT 3": "9PRT3",
                           "PORT 4": "9PRT4",
                           "PORT 5": "9PRT5",
                           "PORT F": "9PRTF"}

    reSetting = re.compile('([^=]+)\s*=\s*(.*)')
    document = Document(documentParam)

    settings = {}

    # Detect relay type SEL-4xx or SEL-3xx
    # Naively check for relay model and decide based on first match
    reRelayModel = re.compile('SEL-([34])([0-9][0-9][A-Z]?)')
    for p in document.paragraphs:
        if reRelayModel.search(p.text):
            rly_type = reRelayModel.search(p.text).group(0)

            logger.debug("Relay type detected: " + rly_type)
            if reRelayModel.search(p.text).group(1) == '3':
                grp_names = sel3XX_names
                rly_family = '3XX'
                break
            elif reRelayModel.search(p.text).group(1) == '4':
                grp_names = sel4XX_names
                rly_family = '4XX'
                break

    for p in document.paragraphs:
        logger.debug('Paragraph style: %s' % p.style.name)
        if p.style.name == 'Heading 1':
            for groupText, reGroup in reGroups:
                if reGroup.match(p.text):
                    if callable(groupText):
                        t = groupText(p.text)
                    else:
                        t = groupText
                    logger.debug(p.style.name + ': "'+p.text+'"')
                    logger.debug("Group:" + t)
                    grp = t
                    break

        if p.style.name == 'Heading 2':
            groupText, reGroup = reGroups[-1]
            if reGroup.match(p.text):
                logger.debug(p.style.name + ': "'+p.text+'"')
                if callable(groupText):
                    t = groupText(p.text)
                else:
                    t = groupText
                grp = t
                logger.debug("Group:" + t)

        if p.style.name in ('SettingLine', 'SettingLineStandard'):
            t = ''.join([r.text for r in filter(setting_run, p.runs)]).strip()
            sm = reSetting.match(t)
            if sm:
                s = [sm.group(1).strip(), sm.group(2).strip()]
                if grp not in settings:
                    settings[grp] = []
                settings[grp].append((s[0], s[1]))
                logger.debug(p.style.name + ': (' + grp + ') "' + s[0] +
                             '", "' + s[1]+'"')
                if reSpecial.search(s[0]):
                    logger.debug('Special character in ' + s[0] + ':' +
                                 ' '.join(map(lambda c: hex(ord(c)),
                                              reSpecial.search(s[0]).group(0))))

    # Some settings are not set individually but are set in ranges in the Word
    # file. For example, PB1 - PB8 = OFF means PB1 = OFF, PB2 = OFF, etc.
    # Looks for integers WITHOUT leading zero. If a leading zero is present, it
    # is treated as part of the text before the number (i.e. PB01-PB09,
    # PB10-PB12 will work but PB01-PB12 will not).
    reRanged = re.compile('([0-9A-Z_]*[0A-Z]+)([1-9][0-9]*)([0-9A-Z_]*)\s*'
                          '[-\u2013]\s*'
                          '\\1([1-9][0-9]*)\\3')
    for grp, settinglist in settings.items():
        rangedSettings, settings[grp] = partition(lambda s: reRanged.match(s[0]), settinglist)
        for rs in rangedSettings:
            p = reRanged.match(rs[0]).groups()
            settings[grp].extend([(p[0] + '%d' % n + p[2], rs[1])
                                  for n in range(int(p[1]), int(p[3])+1)])

    # Some settings are saved in tables rather than in paragraphs
    # Since there are only a few types of settings like this, it is feasible to
    # look for them based on setting names.

    # Protection and Automation Logic sections are only for SEL-4XX relays.
    if rly_family == '4XX':
        # Protection Logic
        # Look for PCTxx, PSVxx, etc. settings
        tLogic = findTable(document.tables, re.compile('(PCT[0-9]{2}PU)|'
                                                       '(PSV[0-9]{2})|'
                                                       '(PMV[0-9]{2})|'
                                                       '(PLT[0-9]{2}S)\s*:=').match)
        if tLogic:
            if len(tLogic.columns) == 2:
                grp = 'PROTECTION LOGIC'
                if grp not in settings:
                    settings[grp] = []
                setList = ['PROTSEL%d' % n for n in range(1, 251)]
                valList = list(filter(lambda s: len(s) > 0,
                                 [''.join([''.join([r.text for r in filter(setting_run, p.runs)]).strip() for p in c.paragraphs]) for c in tLogic.columns[1].cells]))  # NOQA: Needs refactoring
                valList.extend(['']*(250 - len(valList)))
                settings[grp].extend(zip(setList, valList))
            else:
                logger.warning('SEL-4XX Protection Logic table found but '
                               'number of columns is not two.')
        else:
            logger.warning('SEL-4XX Protection Logic table not found.')

        # Automation Logic
        # Look for ASTxx, ASVxx, etc. settings
        tLogic = findTable(document.tables, re.compile('(AST[0-9]{2}IN)|'
                                                       '(ASV[0-9]{3})|'
                                                       'AMV[0-9]{3}\s*:=').match)
        if tLogic:
            if len(tLogic.columns) == 2:
                grp = 'AUTOMATION LOGIC'
                if grp not in settings:
                    settings[grp] = []
                setList = ['AUTO_%d' % n for n in range(1, 101)]
                valList = list(filter(lambda s: len(s) > 0, [''.join([''.join([r.text for r in filter(setting_run, p.runs)]).strip() for p in c.paragraphs]) for c in tLogic.columns[1].cells]))  # NOQA: Needs refactoring
                valList.extend(['']*(100 - len(valList)))
                settings[grp].extend(zip(setList, valList))
            else:
                logger.warning('SEL-4XX Automation Logic table found but '
                               'number of columns is not two.')
        else:
            logger.warning('SEL-4XX Automation Logic table not found.')

    if rly_type == 'SEL-487B':
        # Zone Assignments
        # Look for "Connect Ixx-BZy"
        grp = 'ZONE 1'
        tZoneAssignments = findTable(document.tables,
                                     re.compile('Connect I[01][0-9]-BZ[0-6]').match)
        if tZoneAssignments:
            for r in tZoneAssignments.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        t = ''.join([r.text for r in filter(setting_run, p.runs)]).strip()
                        s = stripall(t.split('='))
                        if len(s) == 2:
                            if grp not in settings:
                                settings[grp] = []
                            settings[grp].append((s[0], s[1]))
                            logger.debug(p.style.name + ': ('+grp+') "' +
                                         s[0]+'", "'+s[1]+'"')
        else:
            logger.warning('SEL-487B zone assignment table not identified')

    # Display Point text and port settings are in tables only for SEL-3xx
    # relays
    if rly_family == '3XX':
        # Display points
        # Look for DP1_1
        tLogic = findTable(document.tables, re.compile('DP1_1\s*=').match)
        grp = 'TEXT'
        if tLogic:
            for r in tLogic.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        t = ''.join([r.text
                                     for r in filter(setting_run, p.runs)]).strip()
                        s = stripall(t.split('='))
                        if len(s) == 2:
                            if grp not in settings:
                                settings[grp] = []
                            settings[grp].append((s[0], s[1]))
                            logger.debug(p.style.name + ': ('+grp+') "' +
                                         s[0] + '", "'+s[1]+'"')
        else:
            logger.warning('SEL-3XX display point table not identified')

        # Port settings
        # Look for EPORT
        tPorts = findTable(document.tables, re.compile('PARITY\s*=').match)
        if tPorts:
            for col in tPorts.columns:
                # First cell in column has a heading the with port number
                # This will not catch columns with headings like "PORTS 1 & 3"
                portMatch = re.match('PORT [F1-4]',
                                     col.cells[0].paragraphs[0].text,
                                     flags=re.I)
                if portMatch:
                    grp = portX(portMatch.group(0))
                    for c in col.cells:
                        for p in c.paragraphs:
                            t = ''.join([r.text for r in filter(setting_run, p.runs)]).strip()
                            s = stripall(t.split('='))
                            if len(s) == 2:
                                if grp not in settings:
                                    settings[grp] = []
                                settings[grp].append((s[0], s[1]))
                                logger.debug(p.style.name + ': ('+grp+') "' +
                                             s[0]+'", "'+s[1]+'"')
        else:
            logger.warning('SEL-3XX Port setting table not identified')

    grp = 'FRONT PANEL'
    # Fill in unset display point settings with blanks so that any that are
    # already set in AcSELerator or Aspen database will be cleared. Settings up
    # through 96 are checked for and added if not present.
    if rly_family == '4XX' and grp in settings:
        setList = set([setName for setName, setVal in settings[grp]])
        for n in range(1, 97):
            if 'DP_ELE%d' % n not in setList:
                settings[grp].append(('DP_ELE%d' % n, ''))


    ###########################
    # Write ASPEN file
    ###########################
    # Read in entire file to a list of lines
    logger.info('Reading Aspen template....')
    with open(aspenTemplateParam, 'r') as f:
        aspen_input = [line.strip() for line in f]

    # The section with the settings is identified by the relay type
    # (e.g. "SEL-351(5,6,7)3").
    # The relay type is identified in the TREQUEST section in the RELAYTYPE
    # field.
    iter_input = iter(aspen_input)
    # Keep a copy of the lines of the file up until the settings start for
    # repeating in the output.
    output_head = []
    for line in iter_input:
        output_head.append(line)
        if line == '[TREQUEST]':
            logger.debug(line)
            break

    relayType = None
    reRelayType = re.compile("RELAYTYPE='(.*)'")
    for line in iter_input:
        output_head.append(line)
        relayType = reRelayType.match(line)
        if relayType:
            logger.debug(line)
            relayType = relayType.group(1)
            break

    if not relayType:
        logger.error('Relay type not detected in AspenDB export file.  Verify '
                     'correct export file was sent to the program.')
        logger.error(__doc__)
        raise SystemExit

    # Choose relay type groups
    if reRelayModel.search(relayType):
        logger.debug("Relay type detected: " +
                     reRelayModel.search(relayType).group(0))
        if reRelayModel.search(relayType).group(1) == '3':
            aspen_names = sel3XX_AspenDBnames
        elif reRelayModel.search(relayType).group(0) == 'SEL-487B':
            aspen_names = sel487B_AspenDBnames
        elif reRelayModel.search(relayType).group(1) == '4':
            aspen_names = sel4XX_AspenDBnames
    else:
        logger.error("Relay type not detected in AspenDB export file.  Verify "
                     "correct export file was sent to the program.")
        logger.error(__doc__)
        raise SystemExit

    inv_aspen_names = invert_dict(aspen_names)

    for line in iter_input:
        output_head.append(line)
        if line == '['+relayType+']':
            logger.debug(line)
            break

    # Remaining lines are settings.  These will be parsed, modified, then
    # re-output for file output
    aspenSettings = {}
    reParseAspenSetting = re.compile("(?P<row>[0-9.]+),"
                                     "'(?P<grp>.+)'="
                                     "'(?P<setting>.*?)',"
                                     "'(?P<range>.*?)',"
                                     "'(?P<value>.*?)',"
                                     "'(?P<comment>.*?)',"
                                     "'.*',"
                                     "'.*',"
                                     "'.*',"
                                     "'.*',"
                                     "'.*'")
    for line in iter_input:
        parseAspenSetting = reParseAspenSetting.match(line)
        if parseAspenSetting:
            if parseAspenSetting.group('grp') in inv_aspen_names:
                grp = parseAspenSetting.group('grp')
                row = parseAspenSetting.group('row')
                setting = parseAspenSetting.group('setting')
                value = parseAspenSetting.group('value')
                setrange = parseAspenSetting.group('range')
                comment = parseAspenSetting.group('comment')
                if grp not in aspenSettings:
                    aspenSettings[grp] = []
                aspenSettings[grp].append({'row': row,
                                           'setting': setting,
                                           'value': value,
                                           'range': setrange,
                                           'comment': comment,
                                           'modified': False})
                logger.debug(aspenSettings[grp][-1])

    # Loop through settings extracted from Word file and try to find a match in
    # the Aspen settings
    logger.info('Matching settings from Word document to Aspen template....')
    grp_missing = set()
    for grp, settinglist in settings.items():
        logger.debug('Processing group '+grp)
        for setting, value in settinglist:
            logger.debug('Looking for setting ' + setting)
            aspen_grp = aspen_names[grp]
            if aspen_grp not in aspenSettings:
                # Skip in loop if group not implemented in Aspen template yet
                if grp not in grp_missing:
                    logger.warning('Group missing from Aspen template: ' + \
                                   grp + ' (' + aspen_grp + ')')
                    grp_missing.add(grp)
                continue
            for aspenSetting in aspenSettings[aspen_grp]:
                # Aspen DB exports setting names all caps, so comparison has to
                # be case insensitive.
                if aspenSetting['setting'].upper().strip() == setting.upper() \
                        and aspenSetting['setting'].upper() != setting.upper():
                    logger.warning('Database setting name with extra space: "'
                                   + aspenSetting['setting'] + '"')
                if aspenSetting['setting'].upper() == setting.upper():
                    logger.debug('setting:' + setting +
                                 'old value:' + aspenSetting['value'] +
                                 'new value:' + value)
                    aspenSetting['value'] = value
                    aspenSetting['modified'] = True
                    # For SEL-351 relays, some of the voltage settings appear
                    # twice due to PTCON = DELTA or WYE global setting.  The
                    # settings that appear twice will be set everywhere they
                    # appear so that the needed setting is set.
                    if not(rly_family == '3XX' and
                           setting in ('27PP', '59PP', '59QP', '59V1P', '27SP',
                                       '59S1', '59S2P')):
                        break

    logger.info('Saving settings to file for import into Aspen....')
    # Loop through final Aspen settings and build string to output to file
    # Only the first three lines and last line of header material is kept,
    # otherwise request data like status, request date, etc. will be
    # overwritten with what is in the template.
    aspen_file_lines = output_head[:3]
    aspen_file_lines.append(output_head[-1])
    for grp, settinglist in aspenSettings.items():
        for setting in settinglist:
            # Only output modified settings
            if setting['modified']:
                aspen_file_lines.append("%s,'%s'='%s','%s','%s','%s',"
                                        "'','','','',''" %
                                        (setting['row'], grp,
                                         setting['setting'], setting['range'],
                                         setting['value'], setting['comment']))
    aspen_file_text = '\n'.join(aspen_file_lines)
    logger.debug(aspen_file_text)

    with codecs.open(aspen_save_file, 'w', encoding='utf-8') as f:
        f.write(aspen_file_text)

    logger.info('Processing special settings for AcSELerator....')
    # Some settings appear as one setting in the setting interface but are
    # represented by multiple settings in the AcSELerator export file.  These
    # must be handled specially to get them to import properly.  The
    # manipulation is done after the ASPEN file is written.

    # DP_ELE (up to 96 settings)
    grp = 'FRONT PANEL'

    if rly_family == '4XX' and rly_type and grp in settings:
        DPSetList = ["DP_ELE", "DP_NAM", "DP_SET", "DP_CLR", "DP_SIZE"]
        DPAnaSetList = ["DP_ELE", "DP_NAM", "DP_SET", "DP_SCA", "DP_CLR", "DP_SIZE"]
        DPSettingList, settings[grp] = partition(lambda s: re.match('DP_ELE', s[0]), settings[grp])
        for DPSetting in DPSettingList:
            n = re.match('DP_ELE([0-9]*)', DPSetting[0]).group(1)
            valList = stripall(stripall(DPSetting[1].split(',')), '"')
            # Identify analog display points by having more settings than usual
            # Other approaches could be used, but this is simple.
            if len(valList) > len(DPSetList):
                setList = [s + n for s in DPAnaSetList]
            else:
                setList = [s + n for s in DPSetList]
            # If not enough settings are provided, extend with blanks
            if len(valList) < len(setList):
                valList.extend(['']*(len(setList) - len(valList)))

            logger.debug(valList)
            settings[grp].extend(zip(setList, valList))

    # SITM (up to 250 settings)
    grp = 'REPORT'
    if rly_family == '4XX' and grp in settings:
        SERSetList = ["SITM", "SNAME", "SSET", "SCLR", "SHMI"]
        SERSettingList, settings[grp] = partition(lambda s: re.match('SITM', s[0]), settings[grp])
        for SERSetting in SERSettingList:
            n = re.match('SITM([0-9]*)', SERSetting[0]).group(1)
            setList = [s + n for s in SERSetList]
            valList = stripall(SERSetting[1].split(','))
            logger.debug(list(zip(setList, valList)))
            settings[grp].extend(zip(setList, valList))

    # SPAQ
    grp = 'REPORT'
    if rly_family == '4XX' and grp in settings:
        SPAQSetting, settings[grp] = partition(lambda s: re.match('SPAQ', s[0]), settings[grp])
        if len(SPAQSetting) > 0:
            valList = stripall(SPAQSetting[0][1].split(','))
            setList = ['SPAQ%d' % n for n in range(1, len(valList) + 1)]
            settings[grp].extend(zip(setList, valList))

    # ERAQ
    grp = 'REPORT'
    if rly_family == '4XX' and grp in settings:
        ERAQSetting, settings[grp] = partition(lambda s: re.match('ERAQ', s[0]), settings[grp])
        if len(ERAQSetting) > 0:
            valList = stripall(ERAQSetting[0][1].split(','))
            setList = ['ERAQ%d' % n for n in range(1, len(valList)+1)]
            settings[grp].extend(zip(setList, valList))

    # SEL-487B Zone Assignments
    #
    # For each terminal Ixx assigned to differential zone BZy, the following
    # settings are made to always assign the current input to the differential
    # zone and to use polarity current (not non-polarity current):
    #    CTxxBZy (Connect Ixx-BZy) = Y
    #    Polarity Ixx-BZz = P  (Polarity / Non-polarity)
    #    IxxBZyV = 1    (Logic equation for dynamic assignment)
    #
    # For all other combinations, CTxxBZy = N.
    #
    # Internal to the relay, these settings are implemented as a list of
    # terminal-differential zone assignments with settings for each assignment
    # as follows:
    #    TBZTn = Ixx   (Terminal being assigned)
    #    TBZBn = BZz   (Bus being assigned to)
    #    TBZPn = P     (Polarity / Non-polarity)
    #    IxxBZyV = 1   (Logic equation for dynamic assignment)
    grp = 'ZONE 1'
    if rly_type == 'SEL-487B' and grp in settings:
        reZoneSetting = re.compile('Connect (?P<TERM>I[012][0-9])'
                                   '-(?P<ZONE>BZ[1-6])', flags=re.I)
        ZONESetting, settings[grp] = partition(lambda s: reZoneSetting.match(s[0]), settings[grp])
        max_zone = 126
        termValList = []
        zoneValList = []
        polValList = []
        for s in ZONESetting:
            m = reZoneSetting.match(s[0])
            if m and s[1] == 'Y':
                termValList.append(m.group('TERM'))
                zoneValList.append(m.group('ZONE'))
                polValList.append(findSetting(settings[grp],
                                              'Polarity %s-%s' %
                                              (m.group('TERM'),
                                               m.group('ZONE'))))
        termSetList = ['TBZT%d' % n for n in range(1, max_zone+1)]
        zoneSetList = ['TBZB%d' % n for n in range(1, max_zone+1)]
        polSetList = ['TBZP%d' % n for n in range(1, max_zone+1)]
        termValList.extend(['']*(max_zone - len(valList)))
        zoneValList.extend(['']*(max_zone - len(valList)))
        polValList.extend(['']*(max_zone - len(valList)))
        settings[grp].extend(zip(termSetList, termValList))
        settings[grp].extend(zip(zoneSetList, zoneValList))
        settings[grp].extend(zip(polSetList, polValList))

    ###########################
    # Write AcSELerator file
    ###########################
    # Make a list of all the settings to output and join together in a string
    # for printing or output to a file.
    #
    # It appears that AcSELerator has some kind of special import mode it uses
    # when there is more than one setting group in a single text file being
    # imported. In this mode, AcSELerator allows at least the some of the
    # SEL-487B settings to be imported with the name set as the description
    # rather than in the internal AcSELerator format. This was found to work
    # for importing SEL-487B polarity settings as "Polarity I01-BZ1" rather
    # than TBZP1.
    # It's not clear how many settings this might apply to, but it seems to
    # work in this particular case and saves having to implement some other
    # logic to convert these settings to the SEL internal format. This quirk
    # didn't seem to work for the zone-bus connections, so the logic to convert
    # to the SEL internal format was added above. Nonetheless, this discovery
    # seems to indicate that some settings might be parsed in this way (i.e.
    # with quotes around the setting name when the name has a space, so the
    # logic below does add quotes around settings names with spaces.
    #

    logger.info('Saving settings to file for import into AcSELerator....')
    sel_file_lines = []
    for grp, settinglist in settings.items():
        sel_file_lines.append('')
        sel_file_lines.append('['+grp_names[grp]+']')
        for setting, value in settinglist:
            sel_file_lines.append((setting if ' ' not in setting else '"'+setting+'"')+',"'+value+'"')

    sel_file_text = '\n'.join(sel_file_lines)
    logger.debug(sel_file_text)

    with codecs.open(sel_save_file, 'w', encoding='utf-8') as f:
        f.write(sel_file_text)

except (SystemExit, KeyboardInterrupt):
    raise
except Exception as e:
    logger.error('Program error', exc_info=True)
finally:
    # Code copied from
    # http://stackoverflow.com/questions/11876618/python-press-any-key-to-exit
    # This uses a Windows-specific library (msvcrt).
    logger.info('DONE.')
    print(u'Press any key to exit...')
    # Assign to a variable just to suppress output. Blocks until key press.
    junk = msvcrt.getch()
