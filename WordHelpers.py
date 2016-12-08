'''WordHelpers.py

These are some utility functions useful for working with Word documents using
python-docx.
'''

from __future__ import print_function, unicode_literals

import docx
from docx.oxml.shared import qn

import operator
import re


# In itertools in Python 3.2+
# https://docs.python.org/3/library/itertools.html#itertools.accumulate
def accumulate(iterable, func=operator.add):
    'Return running totals'
    # accumulate([1,2,3,4,5]) --> 1 3 6 10 15
    # accumulate([1,2,3,4,5], operator.mul) --> 1 2 6 24 120
    it = iter(iterable)
    try:
        total = next(it)
    except StopIteration:
        return
    yield total
    for element in it:
        total = func(total, element)
        yield total


def iter_all_runs(document, *args, **kwargs):
    for p in iter_all_paragraphs(document, *args, **kwargs):
        for r in p.runs:
            yield p, r


def iter_all_paragraphs(document, start=None, end=None):
    ''' Optional parameters start and end cause the itertor to skip until the
        element matches start and quits after element matches end. The range is
        inclusive of start and end.

        This iterator sets the parent of the Paragraph as the parent of the
        element rather than the proxy object for the parent element. It doesn't
        seem to matter since all parent operations are done through the element
        rather than the proxy object.
    '''
    # Uses lxml query to find all, including in tables.
    start_found = start is None
    for p in document._body._element.findall('.//'+qn('w:p')):
        if start_found or p == start:
            start_found = True
            yield docx.text.paragraph.Paragraph(p, p.getparent())
        else:
            continue
        if end is not None and p == end:
            break


def iter_all_paragraphs_old(document):
    ''' This original implementation was discarded in favor of iterating using
        an lxml query that keeps document order and avoids potential issues
        with nested tables.
    '''
    # First scan paragraphs
    for p in document.paragraphs:
        yield p

    # Now scan tables
    for t in document.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    yield p


def delete_paragraph(paragraph):
    p = paragraph._element
    p.clear_content()  # At least remove text
    # Can't remove last paragraph in table cell.
    if isinstance(p.getparent(), docx.oxml.CT_Tc):
        if len(p.getparent().p_lst) < 2:
            return
    p.getparent().remove(p)
    p._p = p._element = None


def delete_run(run):
    r = run._element
    r.getparent().remove(r)
    r._p = r._element = None


def delete_row(row):
    tbl = row.getparent()
    tbl.remove(row)
    # If table no longer has any rows, remove the table.
    if len(tbl.tr_lst) == 0:
        delete_table(tbl)


def delete_table(tbl):
    p = tbl.getparent()
    p.remove(tbl)


def merge_runs(run_list):
    for r in run_list[1:]:
        run_list[0].text += r.text
        delete_run(r)


def find_replace(document, find_text, replace_text, *args, **kwargs):
    find_re = re.compile(find_text)
    for p in iter_all_paragraphs(document, *args, **kwargs):
        if find_re.search(p.text) is not None:
            # Text is in paragraph. Now identify which runs it is in. Keep in
            # mind the text could occur more than once and occur across run
            # breaks.

            # Make list of character index of each run
            run_len = [len(r.text) for r in p.runs]
            run_idx = [0] + list(accumulate(run_len))

            # First pass to merge runs as needed.
            for m in find_re.finditer(p.text):
                start = m.start()
                end = m.end()

                runs_containing = []
                for n, rstart_end in enumerate(zip(run_idx[:-1], run_idx[1:])):
                    rstart, rend = rstart_end
                    if start < rend and end > rstart:
                        runs_containing.append(n)

                # Merge runs if needed.
                if len(runs_containing) > 1:
                    merge_runs([p.runs[n] for n in runs_containing])

            # Second pass to replace text
            # Make list of character index of each run
            run_len = [len(r.text) for r in p.runs]
            run_idx = [0] + list(accumulate(run_len))
            runs_containing = set()
            for m in find_re.finditer(p.text):
                start = m.start()
                end = m.end()

                for n, rstart_end in enumerate(zip(run_idx[:-1], run_idx[1:])):
                    rstart, rend = rstart_end
                    if start < rend and end > rstart:
                        runs_containing.add(n)
                        break
            for n in runs_containing:
                r = p.runs[n]
                rtext = r.text
                newtext = find_re.sub(replace_text, rtext)
                r.text = newtext


def remove_highlighted(document, remove_color, clean_logic_tables=False,
                       *args, **kwargs):
    for p in iter_all_paragraphs(document, *args, **kwargs):
        for r in p.runs:
            if r.font.highlight_color == remove_color:
                r.clear()
                r_p = p
                if len(r_p.text) == 0:
                    delete_paragraph(r_p)
                # If table row is now rendered empty, delete the row
                if isinstance(p._element.getparent(), docx.oxml.CT_Tc):
                    row = p._element.getparent().getparent()
                    row_pars = []
                    for tc in row.tc_lst:
                        for p2 in tc:
                            if isinstance(p2, docx.oxml.CT_P):
                                row_pars.append(docx.text.paragraph.Paragraph(p2, p2.getparent()))  # NOQA: To fix this long line will probably require some refactoring
                    row_text = ''.join([p2.text for p2 in row_pars])
                    if len(row_text) == 0:
                        delete_row(row)
                    # Special rule for Automation Logic & Protection Logic
                    # tables
                    elif clean_logic_tables and \
                        (re.match('[0-9]+\:(NOT 52AA1 *)?$', row_text) or
                         re.match('[0-9]+\:(NOT \(52AA1 OR 52AA2\) *)?$',
                                  row_text)):
                        delete_row(row)

        # There could be empty paragraphs with highlighting applied to
        # paragraph.
        # I had to patch python-docx a little to get this to work.
        if len(p.text) == 0:
            if p.font is not None and p.font.highlight_color == remove_color:
                if p._element.getparent() is not None:
                    delete_paragraph(p)


# http://stackoverflow.com/questions/24965042/python-docx-insertion-point
def get_bookmark_par_element(document, bookmark_name):
    """
    Return the named bookmark parent paragraph element. If no matching
    bookmark is found, the result is '1'. If an error is encountered, '2'
    is returned.
    """
    doc_element = document._body._element
    bookmarks_list = doc_element.findall('.//' + qn('w:bookmarkStart'))
    for bookmark in bookmarks_list:
        name = bookmark.get(qn('w:name'))
        if name == bookmark_name:
            par = bookmark.getparent()
            if not isinstance(par, docx.oxml.CT_P):
                return 2
            else:
                return par
    return 1
